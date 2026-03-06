from docx import Document as DocxDocument

from landrag.ingestion.parsers.html import ParsedDocument, ParsedSection


def extract_docx(file_path: str) -> ParsedDocument:
    doc = DocxDocument(file_path)

    text_parts: list[str] = []
    sections: list[ParsedSection] = []
    current_heading: str | None = None
    current_content: list[str] = []

    for para in doc.paragraphs:
        text_parts.append(para.text)

        if para.style.name.startswith("Heading"):
            if current_heading is not None:
                sections.append(
                    ParsedSection(heading=current_heading, content="\n".join(current_content))
                )
            current_heading = para.text
            current_content = []
        elif current_heading is not None:
            current_content.append(para.text)

    if current_heading is not None:
        sections.append(ParsedSection(heading=current_heading, content="\n".join(current_content)))

    return ParsedDocument(text="\n".join(text_parts), sections=sections)
